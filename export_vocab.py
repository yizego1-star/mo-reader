import json
import sys
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn


GREEN = RGBColor(31, 77, 58)
BLUE = RGBColor(46, 116, 181)
MUTED = RGBColor(104, 119, 111)


def set_font(run, name="Calibri", size=11, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def configure_styles(document):
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
    ]:
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25
        style.paragraph_format.keep_with_next = True


def add_labeled_paragraph(document, label, value, italic=False):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(4)
    label_run = paragraph.add_run(label)
    set_font(label_run, size=10.5, color=MUTED, bold=True)
    value_run = paragraph.add_run(value or "—")
    set_font(value_run, size=10.5, color=GREEN, italic=italic)
    return paragraph


def main():
    if len(sys.argv) != 3:
        raise SystemExit("usage: export_vocab.py input.json output.docx")
    with open(sys.argv[1], "r", encoding="utf-8") as handle:
        payload = json.load(handle)

    document = Document()
    configure_styles(document)

    title = document.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(4)
    title_run = title.add_run("墨读 · 生词本")
    set_font(title_run, size=24, color=GREEN, bold=True)

    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(18)
    subtitle_run = subtitle.add_run(
        f"按加入日期整理 · 共 {len(payload.get('words', []))} 个词 · 导出于 {datetime.now().strftime('%Y-%m-%d')}"
    )
    set_font(subtitle_run, size=10.5, color=MUTED)

    for group in payload.get("groups", []):
        heading = document.add_paragraph(style="Heading 1")
        heading.add_run(group.get("label") or group.get("dateKey") or "未标注日期")
        for item in group.get("words", []):
            word_heading = document.add_paragraph(style="Heading 2")
            word_run = word_heading.add_run(item.get("word") or "—")
            set_font(word_run, size=13, color=GREEN, bold=True)
            pos = item.get("pos") or "词性未记录"
            pos_run = word_heading.add_run(f"  ·  {pos}")
            set_font(pos_run, size=10.5, color=MUTED)
            add_labeled_paragraph(document, "中文意思：", item.get("meaning"))
            if item.get("example"):
                add_labeled_paragraph(document, "例句：", item.get("example"), italic=True)
            if item.get("sourceTitle"):
                add_labeled_paragraph(document, "来源论文：", item.get("sourceTitle"))

    footer = document.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer.add_run("墨读 · 论文精读")
    set_font(footer_run, size=9, color=MUTED)
    document.save(sys.argv[2])


if __name__ == "__main__":
    main()
